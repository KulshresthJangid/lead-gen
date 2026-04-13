from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = Inches(0.65)
section.right_margin = Inches(0.65)
section.top_margin   = Inches(0.5)
section.bottom_margin = Inches(0.5)

# Colour palette
NAVY   = RGBColor(0x0F, 0x2A, 0x4A)
ACCENT = RGBColor(0x1A, 0x73, 0xE8)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
MID    = RGBColor(0x55, 0x55, 0x77)
LIGHT  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ──────────────────────────────────────────────────────────────────

def shade_para(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

def add_run(para, text, bold=False, italic=False, size=10, color=None, font='Calibri'):
    run = para.add_run(text)
    run.font.name   = font
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def bottom_border(para, color='1A73E8', sz='8'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def section_heading(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=10, after=3)
    pPr = p._p.get_or_add_pPr()
    # left blue bar
    pBdr = OxmlElement('w:pBdr')
    lft = OxmlElement('w:left')
    lft.set(qn('w:val'),   'single')
    lft.set(qn('w:sz'),    '18')
    lft.set(qn('w:space'), '6')
    lft.set(qn('w:color'), '1A73E8')
    # bottom rule
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '3')
    bot.set(qn('w:color'), 'D0DCEF')
    pBdr.append(lft)
    pBdr.append(bot)
    pPr.append(pBdr)
    add_run(p, '  ' + text, bold=True, size=11, color=NAVY)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# HEADER BAND
# ══════════════════════════════════════════════════════════════════════════════

name_p = doc.add_paragraph()
shade_para(name_p, '0F2A4A')
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(name_p, before=12, after=2)
add_run(name_p, 'NAMAN SINGHAL', bold=True, size=26, color=LIGHT)

title_p = doc.add_paragraph()
shade_para(title_p, '0F2A4A')
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title_p, before=0, after=12)
add_run(title_p,
        'Senior Backend Engineer  \u00b7  Distributed Systems  \u00b7  Reliability  \u00b7  Event-Driven Architecture',
        size=10.5, color=RGBColor(0xAA, 0xCC, 0xF0))

# Thin accent rule
rule = doc.add_paragraph()
shade_para(rule, '1A73E8')
set_spacing(rule, before=0, after=0, line=4)

# ══════════════════════════════════════════════════════════════════════════════
# IMPACT SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'IMPACT SUMMARY')

impact_items = [
    ('Engineered',  ' event-driven microservices (Kafka + Node.js + Spring Boot) — reduced system downtime by ', '40%', ' and MTTR by ', '28%', '.'),
    ('Optimised',   ' API performance at scale — latency down ', '35%', ', user conversion up ', '12%', '.'),
    ('Automated',   ' CI/CD pipelines with canary releases & rollback, slashing deployment failures by ', '45%', ' and tripling release cadence.', '', ''),
    ('Implemented', ' SLO-based observability that cut production incidents by ', '30%', ' and hardened long-term system resilience.', '', ''),
]

for parts in impact_items:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=2, after=2)
    p.paragraph_format.left_indent = Inches(0.3)
    kw, pre, m1, mid, m2, end = parts
    add_run(p, kw, bold=True, size=10, color=DARK)
    add_run(p, pre, size=10, color=DARK)
    add_run(p, m1, bold=True, size=10, color=ACCENT)
    if mid:
        add_run(p, mid, size=10, color=DARK)
    if m2:
        add_run(p, m2, bold=True, size=10, color=ACCENT)
    if end:
        add_run(p, end, size=10, color=DARK)

# ══════════════════════════════════════════════════════════════════════════════
# CORE TECHNICAL SKILLS
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'CORE TECHNICAL SKILLS')

skills = [
    ('Backend',        'Java (Spring Boot), Node.js, TypeScript, Express.js, REST, GraphQL, gRPC, WebSockets'),
    ('Architecture',   'Microservices, Event-Driven (Kafka), API Design, Concurrency Control, Domain-Driven Design'),
    ('Databases',      'MySQL, PostgreSQL, MongoDB, Redis, Elasticsearch, DynamoDB'),
    ('DevOps & Cloud', 'Kubernetes, Docker, Helm, GitHub Actions, Jenkins, Terraform, AWS'),
    ('Observability',  'Prometheus, Grafana, ELK Stack, SLO Alerting, Automated Test Frameworks'),
]

tbl = doc.add_table(rows=len(skills), cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, items) in enumerate(skills):
    row = tbl.rows[i]
    # Label cell
    c0 = row.cells[0]
    c0.width = Inches(1.6)
    c0_para = c0.paragraphs[0]
    set_spacing(c0_para, before=4, after=4)
    c0_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(c0_para, label, bold=True, size=9.5, color=NAVY)
    # Shade label cell
    tcPr0 = c0._tc.get_or_add_tcPr()
    shd0 = OxmlElement('w:shd')
    shd0.set(qn('w:val'), 'clear')
    shd0.set(qn('w:color'), 'auto')
    shd0.set(qn('w:fill'), 'E8F0FE')
    tcPr0.append(shd0)

    # Items cell
    c1 = row.cells[1]
    c1.width = Inches(5.55)
    c1_para = c1.paragraphs[0]
    set_spacing(c1_para, before=4, after=4)
    c1_para.paragraph_format.left_indent = Inches(0.1)
    add_run(c1_para, items, size=9.5, color=DARK)

# Table borders: only internal hairlines
tblPr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else OxmlElement('w:tblPr')
tblBorders = OxmlElement('w:tblBorders')
for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    el = OxmlElement(f'w:{side}')
    if side.startswith('inside'):
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), 'D0D8E8')
    else:
        el.set(qn('w:val'), 'none')
    tblBorders.append(el)
tblPr.append(tblBorders)

# ══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'PROFESSIONAL EXPERIENCE')

# Job title row
job_p = doc.add_paragraph()
set_spacing(job_p, before=6, after=1)
add_run(job_p, 'Senior Software Engineer', bold=True, size=11.5, color=NAVY)
add_run(job_p, '  \u00b7  Goal-Begins', size=10.5, color=MID)

# Duration with right-aligned tab stop
pPr = job_p._p.get_or_add_pPr()
tabs_el = OxmlElement('w:tabs')
tab_el  = OxmlElement('w:tab')
tab_el.set(qn('w:val'), 'right')
tab_el.set(qn('w:pos'), '9072')   # ~6.3 in from left margin edge
tabs_el.append(tab_el)
pPr.append(tabs_el)
dur = job_p.add_run('\t6+ Years')
dur.font.name   = 'Calibri'
dur.font.size   = Pt(10)
dur.font.italic = True
dur.font.color.rgb = MID

# Subtitle / role context
sub_p = doc.add_paragraph()
set_spacing(sub_p, before=0, after=5)
add_run(sub_p,
        'High-Scale Distributed Systems  \u00b7  Event-Driven Architecture  \u00b7  Engineering Leadership',
        italic=True, size=9.5, color=MID)

exp_bullets = [
    ('High-Scale Systems:',    ' Led backend development for platforms serving millions of daily requests; architected microservices with Kafka-based event streaming.'),
    ('Performance:',           ' Optimised database queries, indexing, and caching strategies — improving throughput and P99 latency by up to 40%.'),
    ('API & Security:',        ' Built production-grade REST and GraphQL APIs with robust authentication, rate-limiting, and contract testing across services.'),
    ('CI/CD & Reliability:',   ' Designed automated pipelines with canary deployments and auto-rollback, cutting failure rates by 45% and tripling release frequency.'),
    ('Observability & SLOs:',  ' Owned SLO monitoring and production incident lifecycle, reducing incidents by 30% through proactive alerting and blameless retrospectives.'),
    ('AI-Augmented Dev:',      ' Embedded AI-assisted workflows for code review, refactoring, and test generation — compressing feature delivery timelines.'),
    ('Leadership & Mentoring:', ' Mentored engineers, established coding standards and quality gates, and drove a culture of continuous improvement across the team.'),
]

for label, detail in exp_bullets:
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=2, after=2)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, label, bold=True, size=10, color=ACCENT)
    add_run(p, detail, size=10, color=DARK)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
spacer = doc.add_paragraph()
set_spacing(spacer, before=10, after=2)
bottom_border(spacer, color='CADAEA', sz='4')

foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(foot, before=3, after=0)
add_run(foot,
        'Open to Senior / Staff Backend Engineering roles  \u00b7  Remote or Hybrid',
        italic=True, size=8.5, color=MID)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('Naman_Singhal_Resume_v3.docx')
print("Done: Naman_Singhal_Resume_v3.docx")
